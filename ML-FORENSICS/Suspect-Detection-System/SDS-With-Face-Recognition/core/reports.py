"""
Detection Report Generator
Generates comprehensive reports of detections and alerts
"""

import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import config


class ReportGenerator:
    """Generate detection reports in various formats."""
    
    def __init__(self, alert_manager, watchlist_manager):
        """
        Initialize report generator.
        
        Args:
            alert_manager: AlertManager instance
            watchlist_manager: WatchlistManager instance
        """
        self.alert_manager = alert_manager
        self.watchlist_manager = watchlist_manager
    
    def generate_session_report(self, session_start, session_end, output_format="docx"):
        """
        Generate a detection session report.
        
        Args:
            session_start: Session start datetime
            session_end: Session end datetime
            output_format: "docx" or "json"
        
        Returns:
            str: Path to generated report
        """
        # Get alerts in time range
        alerts = [a for a in self.alert_manager.alert_history 
                 if session_start <= datetime.fromisoformat(a["timestamp"]) <= session_end]
        
        if output_format == "docx":
            return self._generate_docx_report(alerts, session_start, session_end)
        elif output_format == "json":
            return self._generate_json_report(alerts, session_start, session_end)
        else:
            raise ValueError(f"Unsupported format: {output_format}")
    
    def generate_live_session_report(self, output_format="docx"):
        """
        Generate a report for the current live session (last 24 hours).
        
        Args:
            output_format: "docx" or "json"
        
        Returns:
            str: Path to generated report, or None if no alerts
        """
        session_end = datetime.now()
        session_start = session_end - timedelta(hours=24)
        
        # Get alerts in time range
        alerts = [a for a in self.alert_manager.alert_history 
                 if session_start <= datetime.fromisoformat(a["timestamp"]) <= session_end]
        
        if not alerts:
            return None
        
        return self.generate_session_report(session_start, session_end, output_format)
    
    def _generate_docx_report(self, alerts, session_start, session_end):
        """Generate DOCX format report."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Suspect Detection System - Session Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_heading('Session Information', 1)
        meta_table = doc.add_table(rows=5, cols=2)
        meta_table.style = 'Light Grid Accent 1'
        
        meta_data = [
            ("Report Generated:", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ("Session Start:", session_start.strftime("%Y-%m-%d %H:%M:%S")),
            ("Session End:", session_end.strftime("%Y-%m-%d %H:%M:%S")),
            ("Duration:", str(session_end - session_start)),
            ("Total Alerts:", str(len(alerts)))
        ]
        
        for i, (key, value) in enumerate(meta_data):
            meta_table.rows[i].cells[0].text = key
            meta_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # Summary Statistics
        doc.add_heading('Summary Statistics', 1)
        
        # Count by risk level
        risk_counts = {}
        for alert in alerts:
            risk = alert["risk_level"]
            risk_counts[risk] = risk_counts.get(risk, 0) + 1
        
        # Count by legal status
        status_counts = {}
        for alert in alerts:
            status = alert["legal_status"]
            status_counts[status] = status_counts.get(status, 0) + 1
        
        # Count by camera
        camera_counts = {}
        for alert in alerts:
            camera = alert["camera_id"]
            camera_counts[camera] = camera_counts.get(camera, 0) + 1
        
        # Risk level table
        doc.add_heading('Alerts by Risk Level', 2)
        risk_table = doc.add_table(rows=len(risk_counts)+1, cols=2)
        risk_table.style = 'Light List Accent 1'
        risk_table.rows[0].cells[0].text = "Risk Level"
        risk_table.rows[0].cells[1].text = "Count"
        
        for i, (risk, count) in enumerate(sorted(risk_counts.items()), 1):
            risk_table.rows[i].cells[0].text = risk
            risk_table.rows[i].cells[1].text = str(count)
        
        doc.add_paragraph()
        
        # Camera table
        doc.add_heading('Alerts by Camera', 2)
        camera_table = doc.add_table(rows=len(camera_counts)+1, cols=2)
        camera_table.style = 'Light List Accent 1'
        camera_table.rows[0].cells[0].text = "Camera ID"
        camera_table.rows[0].cells[1].text = "Count"
        
        for i, (camera, count) in enumerate(sorted(camera_counts.items()), 1):
            camera_table.rows[i].cells[0].text = camera
            camera_table.rows[i].cells[1].text = str(count)
        
        doc.add_paragraph()
        doc.add_page_break()
        
        # Detailed Alerts
        doc.add_heading('Detailed Alert Log', 1)
        
        if not alerts:
            doc.add_paragraph("No alerts generated during this session.")
        else:
            for i, alert in enumerate(alerts, 1):
                # Alert header
                heading = doc.add_heading(f"Alert #{i}: {alert['full_name']}", 2)
                
                # Risk level indicator
                risk_para = doc.add_paragraph()
                risk_run = risk_para.add_run(f"  [{alert['risk_level']}]")
                
                # Color code by risk
                if alert['risk_level'] == 'CRITICAL':
                    risk_run.font.color.rgb = RGBColor(220, 38, 38)
                elif alert['risk_level'] == 'HIGH':
                    risk_run.font.color.rgb = RGBColor(234, 88, 12)
                elif alert['risk_level'] == 'MEDIUM':
                    risk_run.font.color.rgb = RGBColor(245, 158, 11)
                
                risk_run.font.bold = True
                
                # Alert details table
                details_table = doc.add_table(rows=12, cols=2)
                details_table.style = 'Light Grid'
                
                details_data = [
                    ("Alert ID:", alert['alert_id']),
                    ("Person ID:", alert['person_id']),
                    ("Case ID:", alert['case_id']),
                    ("Legal Status:", alert['legal_status']),
                    ("Timestamp:", alert['timestamp']),
                    ("Camera:", alert['camera_id']),
                    ("Confidence:", f"{alert['confidence']*100:.1f}%"),
                    ("Occlusion Level:", alert['occlusion_level']),
                    ("Mask Detected:", "Yes" if alert['mask_detected'] else "No"),
                    ("Authorized Agency:", alert['authorized_agency']),
                    ("Evidence Frame:", alert['evidence_frame']),
                    ("Notes:", alert.get('notes', 'N/A'))
                ]
                
                for j, (key, value) in enumerate(details_data):
                    details_table.rows[j].cells[0].text = key
                    details_table.rows[j].cells[1].text = str(value)
                
                doc.add_paragraph()
        
        # Footer
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_para.add_run(f"\n{config.SYSTEM_NAME}\n").bold = True
        footer_para.add_run(f"Version: {config.VERSION}\n")
        footer_para.add_run(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        footer_para.add_run("\nThis is a forensic-grade report. All detections require human verification.\n")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save report
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"detection_report_{timestamp}.docx"
        filepath = os.path.join(config.EVIDENCE_DIR, filename)
        
        doc.save(filepath)
        
        print(f"[REPORT] Generated DOCX report: {filepath}")
        return filepath
    
    def _generate_json_report(self, alerts, session_start, session_end):
        """Generate JSON format report."""
        report = {
            "report_metadata": {
                "generated_at": datetime.now().isoformat(),
                "session_start": session_start.isoformat(),
                "session_end": session_end.isoformat(),
                "duration_seconds": (session_end - session_start).total_seconds(),
                "system_name": config.SYSTEM_NAME,
                "system_version": config.VERSION
            },
            "statistics": {
                "total_alerts": len(alerts),
                "alerts_by_risk": {},
                "alerts_by_status": {},
                "alerts_by_camera": {},
                "unique_persons_detected": len(set(a["person_id"] for a in alerts))
            },
            "alerts": alerts
        }
        
        # Calculate statistics
        for alert in alerts:
            risk = alert["risk_level"]
            status = alert["legal_status"]
            camera = alert["camera_id"]
            
            report["statistics"]["alerts_by_risk"][risk] = \
                report["statistics"]["alerts_by_risk"].get(risk, 0) + 1
            report["statistics"]["alerts_by_status"][status] = \
                report["statistics"]["alerts_by_status"].get(status, 0) + 1
            report["statistics"]["alerts_by_camera"][camera] = \
                report["statistics"]["alerts_by_camera"].get(camera, 0) + 1
        
        # Save report
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"detection_report_{timestamp}.json"
        filepath = os.path.join(config.EVIDENCE_DIR, filename)
        
        with open(filepath, "w") as f:
            json.dump(report, f, indent=2)
        
        print(f"[REPORT] Generated JSON report: {filepath}")
        return filepath
    
    def generate_live_session_report(self):
        """Generate report for current session (from system start)."""
        # Use alert history
        if not self.alert_manager.alert_history:
            print("[REPORT] No alerts to report.")
            return None
        
        first_alert_time = datetime.fromisoformat(self.alert_manager.alert_history[0]["timestamp"])
        last_alert_time = datetime.fromisoformat(self.alert_manager.alert_history[-1]["timestamp"])
        
        return self.generate_session_report(first_alert_time, last_alert_time, "docx")


if __name__ == "__main__":
    print("[INFO] Report generator module loaded.")